"""Generate the synthetic sample corpus (PDF + DOCX + HTML) under ``tests/data/``.

Reproducible source for the small, committed sample documents so the ingest →
chunk → embed → retrieve path is always runnable on a fresh clone (no reliance on
a local/LAN corpus) AND so the spec §7.3 "mixed PDF+DOCX+HTML" claim is actually
satisfied by tracked fixtures. The content is SYNTHETIC, non-PHI, common-knowledge
reference text with two headed sections — enough for Docling's HybridChunker to
emit multiple chunks with heading provenance from each format.

Run:  uv run python scripts/make_sample_corpus.py
"""

from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

DATA_DIR = Path(__file__).resolve().parent.parent / "tests" / "data"
STEM = "sample-clinical-guideline"
TITLE = "Sample Clinical Reference (Synthetic)"

DISCLAIMER = (
    "Synthetic sample document for software testing only. It contains general, "
    "common-knowledge reference text and is not medical advice."
)

# Two logical entries, each a heading + two paragraphs (four paragraphs total).
SECTIONS = [
    (
        "Hand Hygiene in Clinical Settings",
        [
            "Hand hygiene is widely regarded as the single most effective routine "
            "measure for reducing the transmission of common pathogens between "
            "patients and staff. General guidance frames it around clear moments of "
            "care, such as before and after contact with a patient or their "
            "immediate surroundings.",
            "When hands are not visibly soiled, an alcohol-based hand rub is "
            "commonly used for speed and convenience. When hands are visibly soiled, "
            "washing with soap and water is the general reference practice. These are "
            "well-established, non-specific principles rather than situation-specific "
            "instructions.",
        ],
    ),
    (
        "Adult Vital Signs: General Reference Ranges",
        [
            "For a resting adult, commonly cited general reference ranges include a "
            "heart rate of approximately 60 to 100 beats per minute and a "
            "respiratory rate of about 12 to 20 breaths per minute. These figures are "
            "textbook reference values, not thresholds for any individual.",
            "Body temperature is often described around 36.5 to 37.5 degrees Celsius. "
            "Any interpretation for a specific person should come from a qualified "
            "clinician; the values here exist only to provide retrievable, "
            "non-sensitive sample text for testing the retrieval pipeline.",
        ],
    ),
]


def make_pdf() -> Path:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    out = DATA_DIR / f"{STEM}.pdf"
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(out), pagesize=LETTER, title=TITLE)
    flow = [
        Paragraph(TITLE, styles["Title"]),
        Paragraph(DISCLAIMER, styles["Italic"]),
        Spacer(1, 18),
    ]
    for heading, paragraphs in SECTIONS:
        flow.append(Paragraph(heading, styles["Heading1"]))
        for para in paragraphs:
            flow.append(Paragraph(para, styles["BodyText"]))
            flow.append(Spacer(1, 6))
        flow.append(Spacer(1, 12))
    doc.build(flow)
    return out


def make_docx() -> Path:
    from docx import Document as Docx

    out = DATA_DIR / f"{STEM}.docx"
    doc = Docx()
    doc.core_properties.title = TITLE
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(DISCLAIMER, style="Intense Quote")
    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for para in paragraphs:
            doc.add_paragraph(para)
    doc.save(str(out))
    return out


def make_html() -> Path:
    out = DATA_DIR / f"{STEM}.html"
    parts = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="utf-8">',
        f"<title>{escape(TITLE)}</title>",
        "</head>",
        "<body>",
        f"<h1>{escape(TITLE)}</h1>",
        f"<p><em>{escape(DISCLAIMER)}</em></p>",
    ]
    for heading, paragraphs in SECTIONS:
        parts.append(f"<h2>{escape(heading)}</h2>")
        parts.extend(f"<p>{escape(p)}</p>" for p in paragraphs)
    parts += ["</body>", "</html>", ""]
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def build() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    for fn in (make_pdf, make_docx, make_html):
        print(f"Wrote {fn()}")


if __name__ == "__main__":
    build()
